import customtkinter as ctk

class MetadataForm(ctk.CTkFrame):
    def __init__(self, parent, collect_callback):
        super().__init__(parent, fg_color="transparent")
        self.collect_callback = collect_callback
        self._build_form()
        
    def _build_form(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_columnconfigure(3, weight=1)

        # --- HÀNG 1: TỪ NGÀY & ĐẾN NGÀY ---
        # Cụm Từ Ngày
        lbl_from = ctk.CTkLabel(self, text="Từ ngày:")
        lbl_from.grid(row=0, column=0, padx=(5, 2), pady=5, sticky="w")
        
        fr_from = ctk.CTkFrame(self, fg_color="transparent")
        fr_from.grid(row=0, column=1, padx=2, pady=5, sticky="w")
        
        self.ent_from_day = ctk.CTkComboBox(fr_from, values=[str(i) for i in range(1, 31)], width=65)
        self.ent_from_day.pack(side="left", padx=1)
        ctk.CTkLabel(fr_from, text="/").pack(side="left", padx=1)
        self.ent_from_month = ctk.CTkComboBox(fr_from, values=[str(i) for i in range(1, 13)], width=65)
        self.ent_from_month.pack(side="left", padx=1)
        ctk.CTkLabel(fr_from, text="/").pack(side="left", padx=1)
        self.ent_from_year = ctk.CTkComboBox(fr_from, values=["2025", "2026", "2027"], width=85)
        self.ent_from_year.pack(side="left", padx=1)

        # Cụm Đến Ngày
        lbl_to = ctk.CTkLabel(self, text="Đến ngày:")
        lbl_to.grid(row=0, column=2, padx=(20, 2), pady=5, sticky="w")
        
        fr_to = ctk.CTkFrame(self, fg_color="transparent")
        fr_to.grid(row=0, column=3, padx=2, pady=5, sticky="w")
        
        self.ent_to_day = ctk.CTkComboBox(fr_to, values=[str(i) for i in range(1, 31)], width=65)
        self.ent_to_day.pack(side="left", padx=1)
        ctk.CTkLabel(fr_to, text="/").pack(side="left", padx=1)
        self.ent_to_month = ctk.CTkComboBox(fr_to, values=[str(i) for i in range(1, 13)], width=65)
        self.ent_to_month.pack(side="left", padx=1)
        ctk.CTkLabel(fr_to, text="/").pack(side="left", padx=1)
        self.ent_to_year = ctk.CTkComboBox(fr_to, values=["2025", "2026", "2027"], width=85)
        self.ent_to_year.pack(side="left", padx=1)

        # --- HÀNG 2: NHÓM & SỐ NHÂN VIÊN ---
        lbl_group = ctk.CTkLabel(self, text="Nhóm:")
        lbl_group.grid(row=1, column=0, padx=(5, 2), pady=5, sticky="w")
        self.ent_group = ctk.CTkComboBox(self, values=["Số na", "Nhóm 1", "Nhóm 2"], width=235)
        self.ent_group.grid(row=1, column=1, padx=2, pady=5, sticky="w")

        lbl_emp = ctk.CTkLabel(self, text="Số nhân viên:")
        lbl_emp.grid(row=1, column=2, padx=(20, 2), pady=5, sticky="w")
        self.ent_emp_count = ctk.CTkEntry(self, width=85, justify="center")
        self.ent_emp_count.insert(0, "1")
        self.ent_emp_count.grid(row=1, column=3, padx=2, pady=5, sticky="w")

        # --- HÀNG 3: TRANG BỊ & NÚT HÀNH ĐỘNG ---
        lbl_device = ctk.CTkLabel(self, text="Trang bị:")
        lbl_device.grid(row=2, column=0, padx=(5, 2), pady=5, sticky="w")
        self.ent_device = ctk.CTkComboBox(self, values=["Tổ hợp МГК-400ЭМ", "Thiết bị định vị", "Hệ thống radar"], width=235)
        self.ent_device.grid(row=2, column=1, padx=2, pady=5, sticky="w")

        self.btn_action = ctk.CTkButton(self, text="Tạo phiếu BQDP-KSĐK", fg_color="#1f538d", width=180, 
                                        command=lambda: self.collect_callback())
        self.btn_action.grid(row=2, column=3, padx=2, pady=5, sticky="e")

    def get_data(self):
        return {
            "from_date": f"{self.ent_from_day.get()}/{self.ent_from_month.get()}/{self.ent_from_year.get()}",
            "to_date": f"{self.ent_to_day.get()}/{self.ent_to_month.get()}/{self.ent_to_year.get()}",
            "group": self.ent_group.get(),
            "employee_count": self.ent_emp_count.get(),
            "device": self.ent_device.get()
        }
    

import customtkinter as ctk

class WorkTableGrid(ctk.CTkFrame):
    def __init__(self, parent):
        # Đặt màu nền của Frame này làm màu của đường lưới Excel (Gridlines color)
        # Light mode: #d4d4d4 (Xám nhạt Excel) | Dark mode: #3a3a3a
        super().__init__(parent, fg_color=("#d4d4d4", "#3a3a3a"), corner_radius=0)
        self.work_rows_widgets = []
        
        # Thiết lập tỷ lệ co giãn cho các cột (Giữ nguyên cấu trúc logic cũ)
        self.grid_columnconfigure(0, weight=0) # Cột STT/Checkbox cố định rộng
        self.grid_columnconfigure(1, weight=4) # Nội dung công việc rộng nhất
        self.grid_columnconfigure(2, weight=0) # Số người
        self.grid_columnconfigure(3, weight=0) # Thời gian
        self.grid_columnconfigure(4, weight=2) # Phương tiện, dụng cụ
        self.grid_columnconfigure(5, weight=2) # Vật tư
        self.grid_columnconfigure(6, weight=1) # TCKT
        self.grid_columnconfigure(7, weight=1) # Ghi chú
        
        self._build_header()

    def _build_header(self):
        """Xây dựng thanh Header phẳng, vuông vắn chuẩn Excel"""
        headers = ["Mục", "Nội dung công việc", "Số người", "Thời gian", "Phương tiện, dụng cụ", "Vật tư", "TCKT", "Ghi chú"]
        
        # Màu nền Header Excel truyền thống: #f3f3f3 (Light) | #252526 (Dark)
        header_bg = ("#f3f3f3", "#252526")
        
        for col_idx, header_text in enumerate(headers):
            # Tạo một ô Header vuông vức (corner_radius=0)
            # Dùng padx=1, pady=1 trên nền tổng của Frame để lộ ra đường viền ngăn cách mảnh 1px
            cell_frame = ctk.CTkFrame(self, corner_radius=0, fg_color=header_bg, border_width=0, height=30)
            cell_frame.grid(row=0, column=col_idx, sticky="nsew", padx=(1 if col_idx > 0 else 0), pady=(0, 1))
            cell_frame.grid_propagate(False)
            
            lbl = ctk.CTkLabel(
                cell_frame, 
                text=header_text, 
                font=ctk.CTkFont(family="Segoe UI", size=11, weight="bold"), 
                fg_color="transparent",
                text_color=("#333333", "#cccccc")
            )
            lbl.pack(expand=True, fill="both")

    def clear_table(self):
        """Quét sạch các dòng dữ liệu cũ trước khi nạp mới"""
        for row_widgets in self.work_rows_widgets:
            for w in row_widgets.values():
                if w and hasattr(w, "destroy"):
                    w.destroy()
        self.work_rows_widgets.clear()

    def populate_rows(self, structured_data):
        """Đổ dữ liệu mảng vào lưới dạng bảng tính liên mạch"""
        self.clear_table()
        row_counter = 1
        
        # Định nghĩa màu nền cho các ô nhập liệu của Excel
        cell_bg = ("#ffffff", "#1e1e1e")
        
        for item in structured_data:
            if item['type'] != 'table' or not item.get('content'):
                continue
                
            for data_row in item['content'][1:]:
                content_val = data_row[1] if len(data_row) > 1 else ""
                people_val = data_row[2] if len(data_row) > 2 else "0"
                time_val = data_row[3] if len(data_row) > 3 else "15"
                tool_val = data_row[4] if len(data_row) > 4 else ""
                material_val = data_row[5] if len(data_row) > 5 else ""
                tckt_val = data_row[6] if len(data_row) > 6 else ""
                note_val = data_row[7] if len(data_row) > 7 else ""

                # --- Ô 0: CHECKBOX KIÊM CHỈ SỐ MỤC ---
                # Bọc trong ô Frame vuông nền trắng/tối để đồng bộ màu
                cell_chk = ctk.CTkFrame(self, corner_radius=0, fg_color=cell_bg, border_width=0)
                cell_chk.grid(row=row_counter, column=0, sticky="nsew", padx=0, pady=(0, 1))
                
                chk_var = ctk.StringVar(value="on")
                chk = ctk.CTkCheckBox(
                    cell_chk, text=str(row_counter), variable=chk_var, 
                    onvalue="on", offvalue="off", corner_radius=0, width=50,
                    font=ctk.CTkFont(family="Segoe UI", size=11)
                )
                chk.pack(expand=True, padx=5)
                
                # --- Ô 1: NỘI DUNG CÔNG VIỆC (TEXTBOX PHẲNG) ---
                txt_content = ctk.CTkTextbox(
                    self, height=50, border_width=0, corner_radius=0,
                    fg_color=cell_bg, font=ctk.CTkFont(family="Segoe UI", size=11)
                )
                txt_content.insert("1.0", content_val)
                txt_content.grid(row=row_counter, column=1, sticky="nsew", padx=(1, 0), pady=(0, 1))
                
                # --- Ô 2: SỐ NGƯỜI (ENTRY PHẲNG) ---
                ent_people = ctk.CTkEntry(
                    self, width=45, justify="center", border_width=0, corner_radius=0,
                    fg_color=cell_bg, font=ctk.CTkFont(family="Segoe UI", size=11)
                )
                ent_people.insert(0, people_val)
                ent_people.grid(row=row_counter, column=2, sticky="nsew", padx=(1, 0), pady=(0, 1))
                
                # --- Ô 3: THỜI GIAN (COMBOBOX PHẲNG) ---
                cmb_time = ctk.CTkComboBox(
                    self, values=["15", "30", "45", "60", "90", "120"], width=65,
                    border_width=0, corner_radius=0, fg_color=cell_bg, button_color=cell_bg,
                    font=ctk.CTkFont(family="Segoe UI", size=11)
                )
                cmb_time.set(time_val)
                cmb_time.grid(row=row_counter, column=3, sticky="nsew", padx=(1, 0), pady=(0, 1))
                
                # --- Ô 4: PHƯƠNG TIỆN DỤNG CỤ ---
                ent_tool = ctk.CTkEntry(
                    self, border_width=0, corner_radius=0, fg_color=cell_bg,
                    font=ctk.CTkFont(family="Segoe UI", size=11)
                )
                ent_tool.insert(0, tool_val)
                ent_tool.grid(row=row_counter, column=4, sticky="nsew", padx=(1, 0), pady=(0, 1))
                
                # --- Ô 5: VẬT TƯ ---
                ent_material = ctk.CTkEntry(
                    self, border_width=0, corner_radius=0, fg_color=cell_bg,
                    font=ctk.CTkFont(family="Segoe UI", size=11)
                )
                ent_material.insert(0, material_val)
                ent_material.grid(row=row_counter, column=5, sticky="nsew", padx=(1, 0), pady=(0, 1))
                
                # --- Ô 6: TIÊU CHUẨN KỸ THUẬT (TCKT) ---
                ent_tckt = ctk.CTkEntry(
                    self, border_width=0, corner_radius=0, fg_color=cell_bg,
                    font=ctk.CTkFont(family="Segoe UI", size=11)
                )
                ent_tckt.insert(0, tckt_val)
                ent_tckt.grid(row=row_counter, column=6, sticky="nsew", padx=(1, 0), pady=(0, 1))
                
                # --- Ô 7: GHI CHÚ ---
                ent_note = ctk.CTkEntry(
                    self, border_width=0, corner_radius=0, fg_color=cell_bg,
                    font=ctk.CTkFont(family="Segoe UI", size=11)
                )
                ent_note.insert(0, note_val)
                ent_note.grid(row=row_counter, column=7, sticky="nsew", padx=(1, 0), pady=(0, 1))
                
                # Lưu trữ context widget để xuất hoặc đọc ngược lại
                self.work_rows_widgets.append({
                    "cell_chk_frame": cell_chk, # Lưu frame để xóa triệt để khi dọn bảng
                    "check_var": chk_var, "content": txt_content, "people": ent_people,
                    "time": cmb_time, "tool": ent_tool, "material": ent_material,
                    "tckt": ent_tckt, "note": ent_note
                })
                row_counter += 1

    def get_tasks_data(self):
        """Quét dữ liệu từ bảng tính Excel thu về mảng"""
        tasks = []
        for idx, row in enumerate(self.work_rows_widgets):
            if row["check_var"].get() == "on":
                tasks.append({
                    "stt": idx + 1,
                    "content": row["content"].get("1.0", "end-1c").strip(),
                    "people": row["people"].get(),
                    "time": row["time"].get(),
                    "tool": row["tool"].get(),
                    "material": row["material"].get(),
                    "tckt": row["tckt"].get(),
                    "note": row["note"].get()
                })
        return tasks


import os
import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx import Document
from word_processe.config import APP_CONFIG

# Lớp cha điều hướng ứng dụng chính
class MainApplicationView(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title(APP_CONFIG["title"])
        self.geometry(APP_CONFIG["geometry"])
        ctk.set_appearance_mode(APP_CONFIG["theme"]["appearance_mode"])
        
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        self.controller_callback = None
        
        self._init_sidebar()
        self._init_content_area()

    def _init_sidebar(self):
        theme = APP_CONFIG["theme"]
        self.sidebar_frame = ctk.CTkFrame(self, width=240, corner_radius=0, fg_color=theme["sidebar_bg"])
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")
        
        ctk.CTkLabel(self.sidebar_frame, text="★ TƯỜNG VÂN APPS", font=ctk.CTkFont(size=20, weight="bold"), text_color=theme["logo_color"]).grid(row=0, column=0, padx=20, pady=(30, 5))
        ctk.CTkLabel(self.sidebar_frame, text="SYSTEM PORTAL v1.2", font=ctk.CTkFont(size=10), text_color=theme["text_muted"]).grid(row=1, column=0, padx=20, pady=(0, 30))

    def _init_content_area(self):
        self.content_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.content_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=15)
        self.content_frame.grid_rowconfigure(0, weight=1)
        self.content_frame.grid_columnconfigure(0, weight=1)
        
        self.tab_view = ctk.CTkTabview(self.content_frame)
        self.tab_view.grid(row=0, column=0, sticky="nsew")
        
        for tab_cfg in APP_CONFIG["tabs"]:
            tab_obj = self.tab_view.add(tab_cfg["title"])
            if hasattr(self, tab_cfg.get("layout_method", "")):
                getattr(self, tab_cfg["layout_method"])(tab_obj)

    def _build_import_tab_layout(self, parent_tab):
        parent_tab.grid_columnconfigure(0, weight=1)
        parent_tab.grid_rowconfigure(1, weight=1)
        
        # --- 1. ACTION BAR ---
        self.top_action_bar = ctk.CTkFrame(parent_tab, fg_color="transparent")
        self.top_action_bar.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        
        ctk.CTkButton(self.top_action_bar, text="📥 Chọn file Word gốc", command=lambda: self._on_browse_file()).pack(side="left", padx=5)
        self.lbl_file_status = ctk.CTkLabel(self.top_action_bar, text="Chưa nạp tài liệu.", text_color="gray")
        self.lbl_file_status.pack(side="left", padx=15)
        
        ctk.CTkButton(self.top_action_bar, text="⚡ Thu thập dữ liệu Form", fg_color="#2b719e", command=lambda: self.collect_form_data()).pack(side="right", padx=5)
        ctk.CTkButton(self.top_action_bar, text="💾 Xuất Word", fg_color="green", command=lambda: self._export_to_word()).pack(side="right", padx=5)

        # --- 2. FORM AREA ---
        self.form_scroll_frame = ctk.CTkScrollableFrame(parent_tab, border_width=1)
        self.form_scroll_frame.grid(row=1, column=0, sticky="nsew")
        self.form_scroll_frame.grid_columnconfigure(0, weight=1)

        # Nhúng Component Metadata Form
        self.metadata_form = MetadataForm(self.form_scroll_frame, collect_callback=self.collect_form_data)
        self.metadata_form.grid(row=0, column=0, sticky="ew", padx=5, pady=5)
        
        # Nhúng Component Lưới Công Việc
        self.work_table = WorkTableGrid(self.form_scroll_frame)
        self.work_table.grid(row=1, column=0, sticky="ew", padx=5, pady=10)

    def update_document_view(self, file_name, structured_data):
        self.lbl_file_status.configure(text=f"📄 {file_name}", text_color="green")
        self.work_table.populate_rows(structured_data)

    def collect_form_data(self):
        form_data = {
            "metadata": self.metadata_form.get_data(),
            "tasks": self.work_table.get_tasks_data()
        }
        print("--- DỮ LIỆU ĐÃ THU THẬP THÔNG MINH ---")
        print(form_data)
        messagebox.showinfo("Thành công", f"Đã thu thập dữ liệu thông minh từ {len(form_data['tasks'])} hạng mục được chọn!")
        return form_data

    def _export_to_word(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if not file_path: return
        
        current_data = self.collect_form_data()
        doc = Document()
        doc.add_heading(f"KẾ HOẠCH BẢO QUẢN - {current_data['metadata']['device']}", level=1)
        doc.add_paragraph(f"Thời gian: Từ {current_data['metadata']['from_date']} Đến {current_data['metadata']['to_date']}")
        doc.add_paragraph(f"Đơn vị thực hiện: {current_data['metadata']['group']} (Nhân lực: {current_data['metadata']['employee_count']} người)")
        
        table = doc.add_table(rows=1, cols=8)
        hdr_cells = table.rows[0].cells
        headers = ["Mục", "Nội dung công việc", "Số người", "Thời gian", "Phương tiện", "Vật tư", "TCKT", "Ghi chú"]
        for i, text in enumerate(headers):
            hdr_cells[i].text = text
            
        for task in current_data["tasks"]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(task["stt"])
            row_cells[1].text = task["content"]
            row_cells[2].text = task["people"]
            row_cells[3].text = task["time"]
            row_cells[4].text = task["tool"]
            row_cells[5].text = task["material"]
            row_cells[6].text = task["tckt"]
            row_cells[7].text = task["note"]
            
        doc.save(file_path)
        messagebox.showinfo("Thông báo", "Xuất bản tài liệu Word thành công!")

    def _on_browse_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if file_path and self.controller_callback:
            self.controller_callback(file_path)

    def register_controller_handler(self, callback_func):
        self.controller_callback = callback_func


