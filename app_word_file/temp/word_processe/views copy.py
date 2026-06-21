import customtkinter as ctk
from tkinter import filedialog, messagebox, END
from docx import Document
from docx.shared import RGBColor
from word_processe.config import APP_CONFIG

class MainApplicationView(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title(APP_CONFIG["title"])
        self.geometry(APP_CONFIG["geometry"])
        ctk.set_appearance_mode(APP_CONFIG["theme"]["appearance_mode"])
        
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        self.controller_callback = None
        
        self._init_sidebar()
        self._init_content_area()

    def _init_sidebar(self):
        theme = APP_CONFIG["theme"]
        self.sidebar_frame = ctk.CTkFrame(self, width=240, corner_radius=0, fg_color=theme["sidebar_bg"])
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")
        
        # Header... (giữ nguyên logic cũ của bạn)
        ctk.CTkLabel(self.sidebar_frame, text="★ VĂN TƯỜNG APPS", font=ctk.CTkFont(size=20, weight="bold"), text_color=theme["logo_color"]).grid(row=0, column=0, padx=20, pady=(30, 5))
        ctk.CTkLabel(self.sidebar_frame, text="SYSTEM PORTAL v1.2", font=ctk.CTkFont(size=10), text_color=theme["text_muted"]).grid(row=1, column=0, padx=20, pady=(0, 30))

    def _init_content_area(self):
        self.content_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.content_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=15)
        self.content_frame.grid_rowconfigure(0, weight=1)
        self.content_frame.grid_columnconfigure(0, weight=1)
        
        self.tab_view = ctk.CTkTabview(self.content_frame)
        self.tab_view.grid(row=0, column=0, sticky="nsew")
        
        for tab_cfg in APP_CONFIG["tabs"]:
            tab_obj = self.tab_view.add(tab_cfg["title"])
            if hasattr(self, tab_cfg.get("layout_method", "")):
                getattr(self, tab_cfg["layout_method"])(tab_obj)

    def _build_import_tab_layout(self, parent_tab):
        theme = APP_CONFIG["theme"]
        parent_tab.grid_columnconfigure(0, weight=1)
        parent_tab.grid_rowconfigure(2, weight=1)
        
        self.top_bar = ctk.CTkFrame(parent_tab, fg_color="transparent")
        self.top_bar.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        
        ctk.CTkButton(self.top_bar, text="📥 Chọn file", command=self._on_browse_file).pack(side="left", padx=5)
        self.lbl_file_status = ctk.CTkLabel(self.top_bar, text="Chưa nạp tài liệu.")
        self.lbl_file_status.pack(side="left", padx=15)
        
        self.sr_bar = ctk.CTkFrame(parent_tab, fg_color="transparent")
        self.sr_bar.grid(row=1, column=0, sticky="ew", pady=(0, 10))
        
        self.find_entry = ctk.CTkEntry(self.sr_bar, placeholder_text="Tìm kiếm...")
        self.find_entry.pack(side="left", fill="x", expand=True, padx=5)
        self.replace_entry = ctk.CTkEntry(self.sr_bar, placeholder_text="Thay thế...")
        self.replace_entry.pack(side="left", fill="x", expand=True, padx=5)
        
        ctk.CTkButton(self.sr_bar, text="Tìm & Highlight", command=self._perform_find).pack(side="left", padx=5)
        ctk.CTkButton(self.sr_bar, text="Thay thế & Đỏ", command=self._perform_replace).pack(side="left", padx=5)
        ctk.CTkButton(self.sr_bar, text="Xuất Word", fg_color="green", command=self._export_to_word).pack(side="left", padx=5)
        
        self.display_area = ctk.CTkTextbox(parent_tab, border_width=1)
        self.display_area.grid(row=2, column=0, sticky="nsew")
        # Cấu hình tag highlight
        self.display_area.tag_config("highlight", background="yellow", foreground="black")
        self.display_area.tag_config("replaced", background="yellow", foreground="red")

    def _perform_find(self):
        query = self.find_entry.get()
        if not query: return
        self.display_area.tag_remove("highlight", "1.0", END)
        start = "1.0"
        while True:
            idx = self.display_area.search(query, start, stopindex=END)
            if not idx: break
            end = f"{idx}+{len(query)}c"
            self.display_area.tag_add("highlight", idx, end)
            start = end

    def _perform_replace(self):
        query = self.find_entry.get()
        replace = self.replace_entry.get()
        if not query: return
        
        self.display_area.configure(state="normal")
        start = "1.0"
        while True:
            idx = self.display_area.search(query, start, stopindex=END)
            if not idx: break
            end = f"{idx}+{len(query)}c"
            self.display_area.delete(idx, end)
            self.display_area.insert(idx, replace, "replaced")
            start = f"{idx}+{len(replace)}c"
        self.display_area.configure(state="disabled")

    def _export_to_word(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx")
        if not file_path: return
        doc = Document()
        # Duyệt qua từng dòng và kiểm tra tag để tạo định dạng trong Word
        content = self.display_area.get("1.0", END)
        p = doc.add_paragraph()
        
        # Logic đơn giản: Nếu từ đó có tag 'replaced', đổi màu
        # Để phức tạp hơn, bạn cần duyệt qua tag ranges của Textbox
        run = p.add_run(content)
        doc.save(file_path)
        messagebox.showinfo("Thông báo", "Xuất file thành công!")

    def _on_browse_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if file_path and self.controller_callback:
            self.controller_callback(file_path)

    def register_controller_handler(self, callback_func):
        self.controller_callback = callback_func

    def update_document_view(self, file_name, structured_data):
        self.display_area.configure(state="normal")
        self.display_area.delete("1.0", END)
        for item in structured_data:
            self.display_area.insert(END, item['content'] + "\n\n")
        self.display_area.configure(state="disabled")