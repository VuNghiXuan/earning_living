import os
import customtkinter as ctk
from tkinter import filedialog, messagebox, END
from docx import Document
from word_processe.config import APP_CONFIG

class MainApplicationView(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title(APP_CONFIG["title"])
        self.geometry(APP_CONFIG["geometry"])
        ctk.set_appearance_mode(APP_CONFIG["theme"]["appearance_mode"])
        
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        self.controller_callback = None
        
        # Danh sách lưu trữ các widget dòng công việc để lấy dữ liệu thông minh
        self.work_rows_widgets = []
        
        self._init_sidebar()
        self._init_content_area()

    def _init_sidebar(self):
        theme = APP_CONFIG["theme"]
        self.sidebar_frame = ctk.CTkFrame(self, width=240, corner_radius=0, fg_color=theme["sidebar_bg"])
        self.sidebar_frame.grid(row=0, column=0, sticky="nsew")
        
        ctk.CTkLabel(self.sidebar_frame, text="★ TƯỜNG VÂN APPS", font=ctk.CTkFont(size=20, weight="bold"), text_color=theme["logo_color"]).grid(row=0, column=0, padx=20, pady=(30, 5))
        ctk.CTkLabel(self.sidebar_frame, text="SYSTEM PORTAL v1.2", font=ctk.CTkFont(size=10), text_color=theme["text_muted"]).grid(row=1, column=0, padx=20, pady=(0, 30))

    def _init_content_area(self):
        self.content_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.content_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=15)
        self.content_frame.grid_rowconfigure(0, weight=1)
        self.content_frame.grid_columnconfigure(0, weight=1)
        
        self.tab_view = ctk.CTkTabview(self.content_frame)
        self.tab_view.grid(row=0, column=0, sticky="nsew")
        
        for tab_cfg in APP_CONFIG["tabs"]:
            tab_obj = self.tab_view.add(tab_cfg["title"])
            if hasattr(self, tab_cfg.get("layout_method", "")):
                getattr(self, tab_cfg["layout_method"])(tab_obj)

    def _build_import_tab_layout(self, parent_tab):
        parent_tab.grid_columnconfigure(0, weight=1)
        parent_tab.grid_rowconfigure(1, weight=1)
        
        # --- 1. THANH ĐIỀU HƯỚNG CHỌN FILE VÀ XUẤT BÁO CÁO ---
        self.top_action_bar = ctk.CTkFrame(parent_tab, fg_color="transparent")
        self.top_action_bar.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        
        # KHẮC PHỤC: Bọc callback bằng lambda để đảm bảo an toàn context đối tượng
        ctk.CTkButton(self.top_action_bar, text="📥 Chọn file Word gốc", command=lambda: self._on_browse_file()).pack(side="left", padx=5)
        self.lbl_file_status = ctk.CTkLabel(self.top_action_bar, text="Chưa nạp tài liệu.", text_color="gray")
        self.lbl_file_status.pack(side="left", padx=15)
        
        ctk.CTkButton(self.top_action_bar, text="⚡ Thu thập dữ liệu Form", fg_color="#2b719e", command=lambda: self.collect_form_data()).pack(side="right", padx=5)
        ctk.CTkButton(self.top_action_bar, text="💾 Xuất Word", fg_color="green", command=lambda: self._export_to_word()).pack(side="right", padx=5)

        # --- 2. KHU VỰC CHỨA FORM (SCROLLABLE FRAME GIỐNG ẢNH) ---
        self.form_scroll_frame = ctk.CTkScrollableFrame(parent_tab, border_width=1)
        self.form_scroll_frame.grid(row=1, column=0, sticky="nsew")
        self.form_scroll_frame.grid_columnconfigure(0, weight=1)

        # Khởi tạo các nhóm thông tin Metadata phía trên Form
        self._build_metadata_inputs()
        
        # Khởi tạo khung lưới bảng danh sách hạng mục công việc
        self._build_work_table_grid()

    def _build_metadata_inputs(self):
        """Xây dựng khu vực bộ lọc ngày tháng, nhóm, trang bị phía trên giống ảnh"""
        meta_frame = ctk.CTkFrame(self.form_scroll_frame, fg_color="transparent")
        meta_frame.grid(row=0, column=0, sticky="ew", padx=5, pady=5)
        meta_frame.grid_columnconfigure(1, weight=1)

        # Hàng 1: Từ ngày ... Đến ngày
        row1 = ctk.CTkFrame(meta_frame, fg_color="transparent")
        row1.grid(row=0, column=0, columnspan=4, sticky="ew", pady=2)
        
        ctk.CTkLabel(row1, text="Từ ngày").pack(side="left", padx=5)
        self.ent_from_day = ctk.CTkComboBox(row1, values=[str(i) for i in range(1, 32)], width=70)
        self.ent_from_day.pack(side="left", padx=2)
        
        ctk.CTkLabel(row1, text="Tháng").pack(side="left", padx=5)
        self.ent_from_month = ctk.CTkComboBox(row1, values=[str(i) for i in range(1, 13)], width=70)
        self.ent_from_month.pack(side="left", padx=2)
        
        ctk.CTkLabel(row1, text="Năm").pack(side="left", padx=5)
        self.ent_from_year = ctk.CTkComboBox(row1, values=["2025", "2026", "2027"], width=90)
        self.ent_from_year.pack(side="left", padx=2)

        ctk.CTkLabel(row1, text=" |  Đến ngày").pack(side="left", padx=15)
        self.ent_to_day = ctk.CTkComboBox(row1, values=[str(i) for i in range(1, 32)], width=70)
        self.ent_to_day.pack(side="left", padx=2)
        
        ctk.CTkLabel(row1, text="Tháng").pack(side="left", padx=5)
        self.ent_to_month = ctk.CTkComboBox(row1, values=[str(i) for i in range(1, 13)], width=70)
        self.ent_to_month.pack(side="left", padx=2)
        
        ctk.CTkLabel(row1, text="Năm").pack(side="left", padx=5)
        self.ent_to_year = ctk.CTkComboBox(row1, values=["2025", "2026", "2027"], width=90)
        self.ent_to_year.pack(side="left", padx=2)

        # Hàng 2: Nhóm ... Số nhân viên
        row2 = ctk.CTkFrame(meta_frame, fg_color="transparent")
        row2.grid(row=1, column=0, columnspan=4, sticky="ew", pady=2)
        
        ctk.CTkLabel(row2, text="Nhóm").pack(side="left", padx=5)
        self.ent_group = ctk.CTkComboBox(row2, values=["Số na", "Nhóm 1", "Nhóm 2"], width=150)
        self.ent_group.pack(side="left", padx=2)
        
        ctk.CTkLabel(row2, text="Số nhân viên").pack(side="left", padx=20)
        self.ent_emp_count = ctk.CTkEntry(row2, width=80)
        self.ent_emp_count.insert(0, "1")
        self.ent_emp_count.pack(side="left", padx=2)

        # Hàng 3: Trang bị ... Nút chức năng hành động
        row3 = ctk.CTkFrame(meta_frame, fg_color="transparent")
        row3.grid(row=2, column=0, columnspan=4, sticky="ew", pady=2)
        
        ctk.CTkLabel(row3, text="Trang bị").pack(side="left", padx=5)
        self.ent_device = ctk.CTkComboBox(row3, values=["Tổ hợp МГК-400ЭМ", "Thiết bị định vị", "Hệ thống radar"], width=250)
        self.ent_device.pack(side="left", padx=2)
        
        # KHẮC PHỤC: Sử dụng lambda cho nút chức năng chính
        self.btn_action = ctk.CTkButton(row3, text="Tạo phiếu BQDP-KSĐK", fg_color="#1f538d", width=180, command=lambda: self.collect_form_data())
        self.btn_action.pack(side="right", padx=5)

    def _build_work_table_grid(self):
        """Tạo lưới Header cho bảng danh sách hạng mục công việc giống ảnh mẫu"""
        self.table_frame = ctk.CTkFrame(self.form_scroll_frame, fg_color="transparent")
        self.table_frame.grid(row=1, column=0, sticky="ew", padx=5, pady=10)
        
        self.table_frame.grid_columnconfigure(1, weight=4)  # Nội dung công việc rộng nhất
        self.table_frame.grid_columnconfigure(4, weight=2)  # Phương tiện dụng cụ
        self.table_frame.grid_columnconfigure(5, weight=2)  # Vật tư
        self.table_frame.grid_columnconfigure(6, weight=1)  # TCKT
        self.table_frame.grid_columnconfigure(7, weight=1)  # Ghi chú

        headers = ["Mục", "Nội dung công việc", "Số người", "Thời gian", "Phương tiện, dụng cụ", "Vật tư", "TCKT", "Ghi chú"]
        
        for col_idx, header_text in enumerate(headers):
            # Tạo một Frame bọc ngoài để giả lập Border cho Label
            cell_border_frame = ctk.CTkFrame(
                self.table_frame, 
                corner_radius=2, 
                border_width=1, 
                border_color=("#cccccc", "#555555"), # Màu viền cho Light/Dark mode
                fg_color=("#e6e6e6", "#2b2b2b"),
                height=28
            )
            cell_border_frame.grid(row=0, column=col_idx, sticky="nsew", padx=1, pady=1)
            cell_border_frame.grid_propagate(False) # Giữ nguyên chiều cao fix cứng cho đẹp
            
            # Đặt chữ vào trong Frame viền
            lbl = ctk.CTkLabel(
                cell_border_frame, 
                text=header_text, 
                font=ctk.CTkFont(size=12, weight="bold"),
                fg_color="transparent" # Để lộ màu nền của frame bọc ngoài
            )
            lbl.pack(expand=True, fill="both", padx=5)

    def update_document_view(self, file_name, structured_data):
        """Hàm tự động bóc tách dữ liệu mảng 'table' từ file Word để đổ vào Form"""
        for row_widgets in self.work_rows_widgets:
            for w in row_widgets.values():
                if w and hasattr(w, "destroy"): w.destroy()
        self.work_rows_widgets.clear()

        self.lbl_file_status.configure(text=f"📄 {file_name}", text_color="green")
        
        row_counter = 1
        for item in structured_data:
            if item['type'] == 'table':
                table_matrix = item['content']
                if not table_matrix: continue
                
                for data_row in table_matrix[1:]:
                    content_val = data_row[1] if len(data_row) > 1 else ""
                    people_val = data_row[2] if len(data_row) > 2 else "0"
                    time_val = data_row[3] if len(data_row) > 3 else "15"
                    tool_val = data_row[4] if len(data_row) > 4 else ""
                    material_val = data_row[5] if len(data_row) > 5 else ""
                    tckt_val = data_row[6] if len(data_row) > 6 else ""
                    note_val = data_row[7] if len(data_row) > 7 else ""

                    # Cột 0: Checkbox số mục
                    chk_var = ctk.StringVar(value="on")
                    chk = ctk.CTkCheckBox(self.table_frame, text=str(row_counter), variable=chk_var, onvalue="on", offvalue="off", width=50)
                    chk.grid(row=row_counter, column=0, padx=5, pady=2, sticky="w")
                    
                    # Cột 1: Nội dung công việc (Sử dụng Textbox chống tràn chữ)
                    txt_content = ctk.CTkTextbox(self.table_frame, height=60, border_width=1, corner_radius=4)
                    txt_content.insert("1.0", content_val)
                    txt_content.grid(row=row_counter, column=1, padx=2, pady=2, sticky="nsew")
                    
                    # Cột 2: Số người
                    ent_people = ctk.CTkEntry(self.table_frame, width=45, justify="center")
                    ent_people.insert(0, people_val)
                    ent_people.grid(row=row_counter, column=2, padx=2, pady=2)
                    
                    # Cột 3: Thời gian
                    cmb_time = ctk.CTkComboBox(self.table_frame, values=["15", "30", "45", "60", "90", "120"], width=60)
                    cmb_time.set(time_val)
                    cmb_time.grid(row=row_counter, column=3, padx=2, pady=2)
                    
                    # Cột 4: Phương tiện dụng cụ
                    ent_tool = ctk.CTkEntry(self.table_frame, corner_radius=2)
                    ent_tool.insert(0, tool_val)
                    ent_tool.grid(row=row_counter, column=4, padx=2, pady=2, sticky="ew")
                    
                    # Cột 5: Vật tư
                    ent_material = ctk.CTkEntry(self.table_frame, corner_radius=2)
                    ent_material.insert(0, material_val)
                    ent_material.grid(row=row_counter, column=5, padx=2, pady=2, sticky="ew")
                    
                    # Cột 6: TCKT
                    ent_tckt = ctk.CTkEntry(self.table_frame, corner_radius=2)
                    ent_tckt.insert(0, tckt_val)
                    ent_tckt.grid(row=row_counter, column=6, padx=2, pady=2, sticky="ew")
                    
                    # Cột 7: Ghi chú
                    ent_note = ctk.CTkEntry(self.table_frame, corner_radius=2)
                    ent_note.insert(0, note_val)
                    ent_note.grid(row=row_counter, column=7, padx=2, pady=2, sticky="ew")
                    
                    self.work_rows_widgets.append({
                        "check_var": chk_var,
                        "content": txt_content,
                        "people": ent_people,
                        "time": cmb_time,
                        "tool": ent_tool,
                        "material": ent_material,
                        "tckt": ent_tckt,
                        "note": ent_note
                    })
                    row_counter += 1

    def collect_form_data(self):
        """Hàm thông minh quét toàn bộ dữ liệu đang hiển thị trên Form"""
        form_data = {
            "metadata": {
                "from_date": f"{self.ent_from_day.get()}/{self.ent_from_month.get()}/{self.ent_from_year.get()}",
                "to_date": f"{self.ent_to_day.get()}/{self.ent_to_month.get()}/{self.ent_to_year.get()}",
                "group": self.ent_group.get(),
                "employee_count": self.ent_emp_count.get(),
                "device": self.ent_device.get()
            },
            "tasks": []
        }
        
        for idx, row in enumerate(self.work_rows_widgets):
            if row["check_var"].get() == "on":
                task = {
                    "stt": idx + 1,
                    "content": row["content"].get("1.0", "end-1c").strip(),
                    "people": row["people"].get(),
                    "time": row["time"].get(),
                    "tool": row["tool"].get(),
                    "material": row["material"].get(),
                    "tckt": row["tckt"].get(),
                    "note": row["note"].get()
                }
                form_data["tasks"].append(task)
                
        print("--- DỮ LIỆU ĐÃ THU THẬP THÔNG MINH ---")
        print(form_data)
        messagebox.showinfo("Thành công", f"Đã thu thập dữ liệu thông minh từ {len(form_data['tasks'])} hạng mục được chọn!")
        return form_data

    def _perform_find(self):
        pass

    def _perform_replace(self):
        pass

    def _export_to_word(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if not file_path: return
        
        current_data = self.collect_form_data()
        doc = Document()
        doc.add_heading(f"KẾ HOẠCH BẢO QUẢN - {current_data['metadata']['device']}", level=1)
        doc.add_paragraph(f"Thời gian: Từ {current_data['metadata']['from_date']} Đến {current_data['metadata']['to_date']}")
        doc.add_paragraph(f"Đơn vị thực hiện: {current_data['metadata']['group']} (Nhân lực: {current_data['metadata']['employee_count']} người)")
        
        table = doc.add_table(rows=1, cols=8)
        hdr_cells = table.rows[0].cells
        headers = ["Mục", "Nội dung công việc", "Số người", "Thời gian", "Phương tiện", "Vật tư", "TCKT", "Ghi chú"]
        for i, text in enumerate(headers):
            hdr_cells[i].text = text
            
        for task in current_data["tasks"]:
            row_cells = table.add_row().cells
            row_cells[0].text = str(task["stt"])
            row_cells[1].text = task["content"]
            row_cells[2].text = task["people"]
            row_cells[3].text = task["time"]
            row_cells[4].text = task["tool"]
            row_cells[5].text = task["material"]
            row_cells[6].text = task["tckt"]
            row_cells[7].text = task["note"]
            
        doc.save(file_path)
        messagebox.showinfo("Thông báo", "Xuất bản tài liệu Word thành công!")

    def _on_browse_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if file_path and self.controller_callback:
            self.controller_callback(file_path)

    def register_controller_handler(self, callback_func):
        self.controller_callback = callback_func