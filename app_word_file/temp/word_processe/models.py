import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import CT_P, CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

class WordProcessor:
    def __init__(self, file_path=None):
        self.file_path = file_path
        self.doc = None
        if file_path:
            self.read_file(file_path)
        else:
            self.doc = Document()

    def read_file(self, file_path):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Không tìm thấy file tại đường dẫn: {file_path}")
        self.file_path = file_path
        self.doc = Document(file_path)
        return self.doc

    def _clean_text(self, text):
        """Làm sạch các ký tự khoảng trắng đặc biệt (\xa0) tránh lỗi font trên GUI"""
        if not text:
            return ""
        return text.replace('\xa0', ' ').strip()

    def get_structured_content(self):
        """
        Bóc tách cấu trúc file Word bao gồm cả Paragraphs và Tables 
        theo đúng thứ tự xuất hiện một cách tối ưu (O(N)) để đưa lên GUI.
        """
        if not self.doc:
            return []

        structured_data = []
        
        # Duyệt trực tiếp qua các phần tử con trong body XML của tài liệu
        for element in self.doc.element.body:
            
            # 1. XỬ LÝ ĐOẠN VĂN BẢN (PARAGRAPH) - Bao gồm tiêu đề, tên máy, tên phiếu
            if isinstance(element, CT_P):
                p = Paragraph(element, self.doc)
                cleaned_text = self._clean_text(p.text)
                
                if cleaned_text:
                    # Định dạng phân loại Style giúp tầng Giao diện (GUI) biết cách hiển thị font
                    if p.style.name.startswith("Heading") or cleaned_text.isupper() or "PHIẾU SỐ:" in cleaned_text:
                        style_type = "heading"
                    else:
                        style_type = "text"
                        
                    structured_data.append({
                        "type": style_type, 
                        "content": cleaned_text
                    })
            
            # 2. XỬ LÝ BẢNG BIỂU (TABLE) - Danh sách các hạng mục công việc bảo quản
            elif isinstance(element, CT_Tbl):
                table = Table(element, self.doc)
                table_data = []
                
                for row_idx, row in enumerate(table.rows):
                    # Đọc và làm sạch dữ liệu từng ô trong dòng
                    row_data = [self._clean_text(cell.text) for cell in row.cells]
                    
                    # CHỐNG LỆCH CỘT DO GỘP Ô (MERGED CELLS):
                    # Thường dòng tiêu đề thứ 2 (Dụng cụ, Vật tư) sẽ sinh ra dữ liệu lỗi lặp lại.
                    # Nếu dòng có cột đầu (TT) trùng chữ với cột sau, hoặc chứa text tiêu đề lặp, ta bỏ qua dòng rác này trên GUI.
                    if row_idx == 1 and ("P.Tiện" in row_data[0] or "TT" in row_data[0]):
                        continue
                        
                    # Loại bỏ các dòng trống hoàn toàn hoặc dòng rác không có dữ liệu
                    if any(row_data):
                        table_data.append(row_data)
                
                if table_data:
                    structured_data.append({
                        "type": "table", 
                        "content": table_data
                    })
                        
        return structured_data

    def add_paragraph_text(self, text, style=None):
        if not self.doc:
            self.doc = Document()
        self.doc.add_paragraph(text, style=style)

    def save_to_file(self, output_path):
        if not self.doc:
            return False
        try:
            self.doc.save(output_path)
            return True
        except Exception as e:
            print(f"Lỗi khi lưu file: {e}")
            return False