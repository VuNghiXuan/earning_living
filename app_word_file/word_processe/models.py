import os
from docx import Document

class WordProcessor:
    def __init__(self, file_path=None):
        self.file_path = file_path
        self.doc = None
        if file_path:
            self.read_file(file_path)
        else:
            self.doc = Document()

    def read_file(self, file_path):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Không tìm thấy file tại đường dẫn: {file_path}")
        self.file_path = file_path
        self.doc = Document(file_path)
        return self.doc

    def get_structured_content(self):
        """
        Bóc tách cấu trúc file Word bao gồm cả Paragraphs và Tables 
        theo đúng thứ tự xuất hiện để hiển thị lên UI.
        """
        if not self.doc:
            return []

        structured_data = []
        
        # Duyệt qua các thành phần trong body của document
        for element in self.doc.element.body:
            # Nếu là đoạn văn bản (Paragraph)
            if element.tag.endswith('p'):
                for p in self.doc.paragraphs:
                    if p._element == element:
                        if p.text.strip():
                            # Lưu kèm thông tin style để View biết cách render
                            style_type = "heading" if p.style.name.startswith("Heading") else "text"
                            structured_data.append({"type": style_type, "content": p.text})
                        break
            
            # Nếu là bảng biểu (Table)
            elif element.tag.endswith('tbl'):
                for t in self.doc.tables:
                    if t._element == element:
                        table_data = []
                        for row in t.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            table_data.append(row_data)
                        structured_data.append({"type": "table", "content": table_data})
                        break
                        
        return structured_data

    def add_paragraph_text(self, text, style=None):
        if not self.doc:
            self.doc = Document()
        self.doc.add_paragraph(text, style=style)

    def save_to_file(self, output_path):
        if not self.doc:
            return False
        try:
            self.doc.save(output_path)
            return True
        except Exception as e:
            print(f"Lỗi khi lưu file: {e}")
            return False