import sqlite3
import functools
import os
from docx import Document
import sqlite3


class PlanModel:
    def __init__(self, db_path="config/database.db"):
        self.db_path = db_path
        os.makedirs(os.path.dirname(self.db_path), exist_ok=True)
        self.init_db()

    def get_connection(self):
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row # Giữ dòng này để dùng được ["tên_cột"]
        return conn

    def init_db(self):
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("PRAGMA foreign_keys = ON;")
            
            cursor.execute("CREATE TABLE IF NOT EXISTS groups (id INTEGER PRIMARY KEY AUTOINCREMENT, group_name TEXT UNIQUE NOT NULL)")
            cursor.execute("CREATE TABLE IF NOT EXISTS source_files (id INTEGER PRIMARY KEY AUTOINCREMENT, file_name TEXT NOT NULL, file_path TEXT NOT NULL, imported_at DATETIME DEFAULT CURRENT_TIMESTAMP)")
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS devices (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, group_id INTEGER NOT NULL, source_file_id INTEGER, device_name TEXT UNIQUE NOT NULL,
                    FOREIGN KEY (group_id) REFERENCES groups(id) ON DELETE CASCADE, FOREIGN KEY (source_file_id) REFERENCES source_files(id) ON DELETE SET NULL
                )
            """)
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS maintenance_cycles (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, 
                    device_id INTEGER NOT NULL, 
                    cycle_code TEXT NOT NULL, 
                    cycle_name TEXT, 
                    
                    FOREIGN KEY (device_id) REFERENCES devices(id) ON DELETE CASCADE
                )
            """)                        
            # Cập nhạt bảng định mức
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS master_tasks (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, 
                    cycle_id INTEGER NOT NULL, 
                    tt TEXT NOT NULL, 
                    task_name TEXT NOT NULL,
                    norm_workers INTEGER DEFAULT 1, 
                    norm_minutes INTEGER DEFAULT 0, 
                    tool TEXT, 
                    material TEXT, 
                    tckt TEXT, 
                    result TEXT,
                    order_index INTEGER DEFAULT 0, -- THÊM DÒNG NÀY
                    FOREIGN KEY (cycle_id) REFERENCES maintenance_cycles(id) ON DELETE CASCADE
                )
            """)

            # Thêm bảng system_settings vào hàm init_db của bạn:
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS system_settings (
                    key TEXT PRIMARY KEY,
                    value TEXT
                )
            """)

    # Thêm hàm bắt lỗi:
    def handle_db_error(func):
        @functools.wraps(func)
        def wrapper(self, *args, **kwargs): # Thêm self vào đây
            try:
                return func(self, *args, **kwargs)
            except sqlite3.OperationalError as e:
                print(f"[ERROR DETECTED] trong hàm {func.__name__}: {e}")
                raise e
        return wrapper

    
    

    @handle_db_error
    def import_from_word_data(self, parsed_data, target_group_name, file_name, file_path):
        if not parsed_data:
            return
        
        if not target_group_name or target_group_name.strip() == "":
            target_group_name = "Chưa phân loại"
            
        with self.get_connection() as conn:
            cursor = conn.cursor()
            
            # 1. Đảm bảo nhóm tồn tại
            cursor.execute("INSERT OR IGNORE INTO groups (group_name) VALUES (?)", (target_group_name,))
            cursor.execute("SELECT id FROM groups WHERE group_name = ?", (target_group_name,))
            group_id = cursor.fetchone()["id"]
            
            # 2. Ghi nhận file nguồn
            cursor.execute("INSERT INTO source_files (file_name, file_path) VALUES (?, ?)", (file_name, file_path))
            source_file_id = cursor.lastrowid
            
            # 3. Nạp thiết bị và tác vụ
            for dev_obj in parsed_data:
                d_name = dev_obj.device_name
                
                # Insert thiết bị
                cursor.execute("INSERT OR IGNORE INTO devices (group_id, source_file_id, device_name) VALUES (?, ?, ?)", 
                               (group_id, source_file_id, d_name))
                cursor.execute("SELECT id FROM devices WHERE device_name = ?", (d_name,))
                device_id = cursor.fetchone()["id"]
                
                # Duyệt qua các chu kỳ (cycles là dict: {"Phiếu 01": [tasks]})
                for c_code, tasks in dev_obj.cycles.items():
                    # Kiểm tra chu kỳ đã tồn tại chưa
                    cursor.execute("SELECT id FROM maintenance_cycles WHERE device_id = ? AND cycle_code = ?", (device_id, c_code))
                    cycle_row = cursor.fetchone()
                    
                    if cycle_row:
                        cycle_id = cycle_row["id"]
                        # Xóa các task cũ để nạp mới (cập nhật nội dung)
                        cursor.execute("DELETE FROM master_tasks WHERE cycle_id = ?", (cycle_id,))
                    else:
                        
                        cursor.execute("""
                            INSERT INTO maintenance_cycles (device_id, cycle_code, cycle_name) 
                            VALUES (?, ?, ?)
                        """, (device_id, c_code, c_code)) 
                        cycle_id = cursor.lastrowid
                    
                    # Insert các đầu việc
                    # Trong vòng lặp for t in tasks (nạp các đầu việc):
                    for t in tasks:
                        cursor.execute("""
                            INSERT INTO master_tasks (cycle_id, tt, task_name, norm_workers, norm_minutes, tool, material, tckt, result, order_index)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """, (cycle_id, t.tt, t.name, t.workers, t.minutes, t.tool, t.material, t.tckt, t.result, t.order_index)) # Thêm t.order_index
            
            conn.commit()

    # Code cho system_settings
    def save_setting(self, key, value):
        with self.get_connection() as conn:
            conn.execute("INSERT OR REPLACE INTO system_settings (key, value) VALUES (?, ?)", (key, str(value)))

    # def get_setting(self, key, default=None):
    #     with self.get_connection() as conn:
    #         row = conn.execute("SELECT value FROM system_settings WHERE key = ?", (key,)).fetchone()
    #         return row["value"] if row else default
        
    # Code cho nhóm sử dụng máy
    @handle_db_error
    def move_device_to_group(self, device_name, new_group_name):
        with self.get_connection() as conn:
            cursor = conn.cursor()
            # Lấy ID của nhóm đích
            cursor.execute("SELECT id FROM groups WHERE group_name = ?", (new_group_name,))
            group_row = cursor.fetchone()
            if group_row:
                new_group_id = group_row["id"]
                # Cập nhật group_id cho máy
                cursor.execute("UPDATE devices SET group_id = ? WHERE device_name = ?", (new_group_id, device_name))
                conn.commit()

    @handle_db_error
    def generate_daily_plan_word(self, save_path, selected_tasks, device_name, cycle_info, plan_date_str):
        doc = Document()
        doc.add_heading('KẾ HOẠCH BẢO DƯỠNG TRANG BỊ TRONG NGÀY', level=1)
        doc.add_paragraph(f"Ngày lập thực hiện: {plan_date_str}")
        doc.add_paragraph(f"Tên trang bị: {device_name}")
        doc.add_paragraph(f"Nội dung thực hiện: {cycle_info}")
        
        # 1. Tăng số cột lên 8
        table = doc.add_table(rows=1, cols=8) 
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        
        # 2. Thêm tiêu đề 'Kết quả' vào headers
        headers = ['STT', 'Nội dung công việc', 'Số người', 'Thời gian (Phút)', 'Dụng cụ', 'Vật tư', 'TCKT', 'Kết quả']
        for idx, text in enumerate(headers):
            hdr_cells[idx].text = text
            
        for task in selected_tasks:
            row_cells = table.add_row().cells
            row_cells[0].text = str(task.get("tt", ""))
            row_cells[1].text = str(task.get("name", ""))
            row_cells[2].text = str(task.get("workers", ""))
            row_cells[3].text = str(task.get("minutes", ""))
            row_cells[4].text = str(task.get("tool", ""))
            row_cells[5].text = str(task.get("material", ""))
            row_cells[6].text = str(task.get("tckt", ""))
            # 3. Nạp dữ liệu vào cột thứ 8 (index 7)
            row_cells[7].text = str(task.get("result", "")) 
            
        doc.save(save_path)

    # Cập nhật dữ liệu cho danh mục định mức
    @handle_db_error
    def get_all_norms(self):
        query = """
            SELECT d.device_name, mc.cycle_code,  mt.tt, mt.task_name, 
                mt.norm_workers, mt.norm_minutes, mt.tool, mt.material, mt.tckt, mt.result
            FROM master_tasks mt
            JOIN maintenance_cycles mc ON mt.cycle_id = mc.id
            JOIN devices d ON mc.device_id = d.id
            ORDER BY d.device_name, mc.cycle_code, CAST(mt.tt AS REAL) ASC
        """
        with self.get_connection() as conn:
            return [list(row) for row in conn.execute(query).fetchall()]

    @handle_db_error    
    def get_tasks(self, group_name, device_name, cycle_code):
        """
        Truy vấn danh sách công việc, sắp xếp theo order_index.
        """
        with self.get_connection() as conn:
            query = """
                SELECT t.* FROM master_tasks t
                JOIN maintenance_cycles c ON t.cycle_id = c.id
                JOIN devices d ON c.device_id = d.id
                JOIN groups g ON d.group_id = g.id
                WHERE g.group_name = ? AND d.device_name = ? AND c.cycle_code = ?
                ORDER BY t.order_index ASC -- QUAN TRỌNG: Sắp xếp theo order_index
            """
            return [dict(row) for row in conn.execute(query, (group_name, device_name, cycle_code)).fetchall()]
    
    @handle_db_error    
    def get_devices_by_group(self, group_name):
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT d.device_name FROM devices d
                JOIN groups g ON d.group_id = g.id
                WHERE g.group_name = ?
            """, (group_name,))
            return [row["device_name"] for row in cursor.fetchall()]

    # @handle_db_error
    # def get_all_group_names(self):
    #     with self.get_connection() as conn:
    #         cursor = conn.cursor()
    #         cursor.execute("SELECT group_name FROM groups")
    #         return [row["group_name"] for row in cursor.fetchall()]
    
    # @handle_db_error
    # def update_all_norms(self, mapped_data_list):
    #     if not mapped_data_list: return False

    #     with self.get_connection() as conn:
    #         conn.row_factory = None
    #         cursor = conn.cursor()
    #         try:
    #             # 1. BẮT ĐẦU TRANSACTION: Xóa sạch dữ liệu cũ trong 3 bảng liên quan
    #             # Việc xóa theo thứ tự ngược lại (Tasks -> Cycles -> Devices) để tránh lỗi Foreign Key
    #             cursor.execute("BEGIN TRANSACTION")
                
    #             cursor.execute("DELETE FROM master_tasks")
    #             cursor.execute("DELETE FROM maintenance_cycles")
    #             cursor.execute("DELETE FROM devices")
    #             # Nếu muốn xóa luôn cả groups, hãy thêm lệnh xóa ở đây
                
    #             print("[DEBUG] Đã xóa toàn bộ dữ liệu cũ trong DB.")

    #             # 2. CHÈN DỮ LIỆU MỚI (Tái tạo hệ thống)
    #             # Dùng dict để tracking các ID đã tạo để tránh trùng lặp khi Insert
    #             device_cache = {} # { "Tên thiết bị": id }
    #             cycle_cache = {}  # { ("Tên thiết bị", "Code"): id }

    #             for item in mapped_data_list:
    #                 d_name = str(item['device_name']).strip()
    #                 c_code = str(item['cycle_code']).strip()

    #                 # A. Xử lý Thiết bị
    #                 if d_name not in device_cache:
    #                     cursor.execute("INSERT INTO devices (group_id, device_name) VALUES (?, ?)", (1, d_name))
    #                     device_cache[d_name] = cursor.lastrowid
    #                 device_id = device_cache[d_name]

    #                 # B. Xử lý Chu kỳ
    #                 cycle_key = (d_name, c_code)
    #                 if cycle_key not in cycle_cache:
    #                     cursor.execute("INSERT INTO maintenance_cycles (device_id, cycle_code, cycle_name) VALUES (?, ?, ?)", 
    #                                    (device_id, c_code, c_code))
    #                     cycle_cache[cycle_key] = cursor.lastrowid
    #                 cycle_id = cycle_cache[cycle_key]

    #                 # C. Xử lý Task (Vì bảng đã xóa sạch nên cứ INSERT thôi)
    #                 cursor.execute("""
    #                     INSERT INTO master_tasks (cycle_id, tt, task_name, norm_workers, norm_minutes, tool, material, tckt, result)
    #                     VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    #                 """, (cycle_id, item['tt'], item['task_name'], item['workers'], item['minutes'], 
    #                       item['tool'], item['material'], item['tckt'], item['result'], item.get('order_index', 0)))

    #             conn.commit()
    #             print("[DEBUG] Đã cập nhật lại toàn bộ DB từ giao diện.")
    #             return True
    #         except Exception as e:
    #             conn.rollback()
    #             print(f"[ERROR] Quá trình cập nhật thất bại, đã rollback: {e}")
    #             return False
    
    @handle_db_error
    def update_all_norms(self, mapped_data_list):
        if not mapped_data_list: return False

        with self.get_connection() as conn:
            conn.row_factory = None
            cursor = conn.cursor()
            try:
                cursor.execute("BEGIN TRANSACTION")
                # Xóa dữ liệu cũ
                cursor.execute("DELETE FROM master_tasks")
                cursor.execute("DELETE FROM maintenance_cycles")
                cursor.execute("DELETE FROM devices")

                device_cache = {} 
                cycle_cache = {} 

                for item in mapped_data_list:
                    # Sử dụng khóa 'DEVICE' và 'CYCLE' khớp với NORMS_TABLE_CONFIG
                    d_name = str(item.get('DEVICE', 'Chưa đặt tên')).strip() 
                    c_code = str(item.get('CYCLE', 'Phiếu mặc định')).strip()

                    # A. Xử lý Thiết bị
                    if d_name not in device_cache:
                        cursor.execute("INSERT INTO devices (group_id, device_name) VALUES (?, ?)", (1, d_name))
                        device_cache[d_name] = cursor.lastrowid
                    device_id = device_cache[d_name]

                    # B. Xử lý Chu kỳ
                    cycle_key = (d_name, c_code)
                    if cycle_key not in cycle_cache:
                        cursor.execute("INSERT INTO maintenance_cycles (device_id, cycle_code, cycle_name) VALUES (?, ?, ?)", 
                                       (device_id, c_code, c_code))
                        cycle_cache[cycle_key] = cursor.lastrowid
                    cycle_id = cycle_cache[cycle_key]

                    # C. Xử lý Task (Dùng .get() để lấy giá trị an toàn)
                    cursor.execute("""
                        INSERT INTO master_tasks (cycle_id, tt, task_name, norm_workers, norm_minutes, tool, material, tckt, result, order_index)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (cycle_id, item.get('TT'), item.get('TASK'), item.get('WORKERS', 0), item.get('MINUTES', 0), 
                          item.get('TOOL'), item.get('MATERIAL'), item.get('TCKT'), item.get('RESULT'), item.get('ORDER', 0)))

                conn.commit()
                print("[DEBUG] Đã cập nhật lại toàn bộ DB thành công.")
                return True
            except Exception as e:
                conn.rollback()
                print(f"[ERROR] Quá trình cập nhật thất bại, đã rollback: {e}")
                return False
            
    # Lấy Phiếu số từ DB đưa vào setting_bar
    # @handle_db_error
    # def get_all_phieu_names(self):
    #     """
    #     Lấy danh sách tên phiếu (cycle_code) từ bảng maintenance_cycles.
    #     """
    #     try:
    #         # Gọi phương thức get_connection() để lấy kết nối DB
    #         conn = self.get_connection()
    #         cursor = conn.cursor()
            
    #         # Truy vấn lấy tên phiếu
    #         query = "SELECT DISTINCT cycle_code FROM maintenance_cycles WHERE cycle_code IS NOT NULL AND cycle_code != ''"
            
    #         cursor.execute(query)
    #         results = [row[0] for row in cursor.fetchall()]
            
    #         # Đóng kết nối nếu cần thiết (tùy vào cách cấu trúc DB của bạn)
    #         # conn.close() 
            
    #         return results
    #     except Exception as e:
    #         print(f"[ERROR] Không thể lấy danh sách phiếu từ maintenance_cycles: {e}")
    #         return []
    
    # Lọc máy (trang bị) khi biết phiếu số (01,02,03,04) và nhóm
    # @handle_db_error
    # def get_devices_by_group_and_cycle(self, group_name, cycle_code):
    #     with self.get_connection() as conn:
    #         cursor = conn.cursor()
    #         # Lọc theo Nhóm VÀ có chứa Phiếu công nghệ (maintenance_cycles)
    #         query = """
    #             SELECT DISTINCT d.device_name 
    #             FROM devices d
    #             JOIN groups g ON d.group_id = g.id
    #             JOIN maintenance_cycles mc ON d.id = mc.device_id
    #             WHERE g.group_name = ? AND mc.cycle_code = ?
    #         """
    #         cursor.execute(query, (group_name, cycle_code))
    #         return [row["device_name"] for row in cursor.fetchall()]
    
    # Các hàm thêm xoá tên nhóm:
    def add_new_group(self, group_name):        
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO groups (group_name) VALUES (?)", (group_name.strip(),))
            conn.commit()
            
    # def add_group(self, name):
    #     with self.get_connection() as conn:
    #         conn.execute("INSERT INTO groups (group_name) VALUES (?)", (name,))
    #         conn.commit()

    def rename_group(self, old_name, new_name):
        with self.get_connection() as conn:
            cursor = conn.cursor()
            # 1. Lấy ID của nhóm cũ
            cursor.execute("SELECT id FROM groups WHERE group_name = ?", (old_name,))
            row = cursor.fetchone()
            if row:
                group_id = row["id"]
                # 2. Đổi tên nhóm trong bảng groups
                cursor.execute("UPDATE groups SET group_name = ? WHERE id = ?", (new_name, group_id))
                conn.commit()

    def delete_group(self, name):
        with self.get_connection() as conn:
            cursor = conn.cursor()
            # 1. Lấy ID của nhóm cần xóa
            cursor.execute("SELECT id FROM groups WHERE group_name = ?", (name,))
            row = cursor.fetchone()
            if row:
                group_id = row["id"]
                # 2. Chuyển các máy trong nhóm đó về nhóm "Chưa phân loại"
                # Giả định ID của nhóm "Chưa phân loại" là 1 (hoặc bạn cần query nó ra)
                cursor.execute("UPDATE devices SET group_id = 1 WHERE group_id = ?", (group_id,))
                # 3. Xóa nhóm
                cursor.execute("DELETE FROM groups WHERE id = ?", (group_id,))
                conn.commit()

    # Cập nhật bộ lọc thông tin
    def get_all_phieu_names(self):
        with self.get_connection() as conn:
            rows = conn.execute("SELECT DISTINCT cycle_code FROM maintenance_cycles").fetchall()
            return [r["cycle_code"] for r in rows]

    def get_all_group_names(self):
        # Đảm bảo lệnh truy vấn lấy đúng dữ liệu hiện tại
        with self.get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT group_name FROM groups") # Kiểm tra lại tên bảng của bạn
            return [row[0] for row in cursor.fetchall()]

    def get_devices_by_group_and_cycle(self, group_name, cycle_code):
        if group_name == "Chưa phân loại": return []
        with self.get_connection() as conn:
            cursor = conn.cursor()
            # Đảm bảo join đúng bảng
            query = """
                SELECT d.device_name 
                FROM devices d
                JOIN groups g ON d.group_id = g.id
                JOIN maintenance_cycles mc ON d.id = mc.device_id
                WHERE g.group_name = ? AND mc.cycle_code = ?
            """
            cursor.execute(query, (group_name, cycle_code))
            return [row[0] for row in cursor.fetchall()]

    def get_setting(self, key, default="0"):
        with self.get_connection() as conn:
            row = conn.execute("SELECT value FROM system_settings WHERE key = ?", (key,)).fetchone()
            return row["value"] if row else default
    
    # Thay truy xuất ngược, lấy phiếu bảo dưỡng 01, 02 -> thuộc nhóm nào
    def get_groups_by_cycle(self, cycle_code):
        query = """
            SELECT DISTINCT g.group_name FROM groups g
            JOIN devices d ON g.id = d.group_id
            JOIN maintenance_cycles mc ON d.id = mc.device_id
            WHERE mc.cycle_code = ?
        """
        with self.get_connection() as conn:
            return [r[0] for r in conn.execute(query, (cycle_code,)).fetchall()]

    def get_cycles_by_group(self, group_name):
        query = """
            SELECT DISTINCT mc.cycle_code FROM maintenance_cycles mc
            JOIN devices d ON mc.device_id = d.id
            JOIN groups g ON d.group_id = g.id
            WHERE g.group_name = ?
        """
        with self.get_connection() as conn:
            return [r[0] for r in conn.execute(query, (group_name,)).fetchall()]
    
    def get_devices_by_filter(self, group_name=None, cycle_code=None):
        """Hàm lấy danh sách máy dựa trên nhóm hoặc phiếu hoặc cả hai"""
        query = """
            SELECT DISTINCT d.device_name 
            FROM devices d
            JOIN groups g ON d.group_id = g.id
            LEFT JOIN maintenance_cycles mc ON d.id = mc.device_id
            WHERE 1=1
        """
        params = []
        
        if group_name and group_name != "-- Chọn nhóm --" and group_name != "Chưa phân loại":
            query += " AND g.group_name = ?"
            params.append(group_name)
            
        if cycle_code and cycle_code != "-- Chọn phiếu BQDP --":
            query += " AND mc.cycle_code = ?"
            params.append(cycle_code)
            
        with self.get_connection() as conn:
            cursor = conn.execute(query, params)
            return [r[0] for r in cursor.fetchall()]
    
    # Đưa định mức lên bảng kế hoạch
    def get_tasks_by_filter(self, group_name, phieu_name):
        # Chúng ta JOIN từ bảng master_tasks ngược về các bảng liên quan
        query = """
            SELECT
                d.device_name, 
                t.tt, 
                t.task_name, 
                t.norm_workers AS std_people, 
                t.norm_minutes AS std_time,
                t.tool, t.material, t.tckt, t.result
            FROM master_tasks t
            JOIN maintenance_cycles mc ON t.cycle_id = mc.id
            JOIN devices d ON mc.device_id = d.id
            JOIN groups g ON d.group_id = g.id
            WHERE 1=1
        """
        params = []
        
        if phieu_name and phieu_name != "-- Chọn phiếu BQDP --":
            query += " AND mc.cycle_code = ?"
            params.append(phieu_name)
            
        if group_name and group_name != "Tất cả các nhóm":
            query += " AND g.group_name = ?"
            params.append(group_name)
            
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute(query, params)
                columns = [col[0] for col in cursor.description]
                return [dict(zip(columns, row)) for row in cursor.fetchall()]
        except Exception as e:
            print(f"[ERROR] Database query failed: {e}")
            return []
    
    def get_norms_by_filter(self, filter_results):
        """
        Lấy dữ liệu từ bảng master_tasks thông qua các bảng liên quan.
        Có bổ sung debug để truy vết lỗi cấu trúc dữ liệu.
        """
        # 1. DEBUG INPUT
        print(f"[DEBUG MODEL] Bắt đầu truy vấn với filter_results: {filter_results}")
        
        trang_bi_list = filter_results.get('trang_bi', [])
        if not trang_bi_list:
            print("[DEBUG MODEL] CẢNH BÁO: trang_bi_list bị trống!")
            return []
        
        ten_may_list = [item['ten_may'] for item in trang_bi_list if 'ten_may' in item]
        print(f"[DEBUG MODEL] Danh sách máy lọc được: {ten_may_list}")
        
        if not ten_may_list:
            print("[DEBUG MODEL] CẢNH BÁO: Không có tên máy nào hợp lệ trong trang_bi_list.")
            return []
            
        nhom = filter_results.get('nhom_thuc_hien', '')
        print(f"[DEBUG MODEL] Nhóm lọc được: {nhom}")
        
        # 2. Xây dựng SQL
        placeholders = ', '.join(['?'] * len(ten_may_list))
        query = f"""
            SELECT 
                d.device_name, t.tt, t.task_name, 
                t.norm_workers, t.norm_minutes,
                t.tool, t.material, t.tckt, t.result
            FROM master_tasks t
            JOIN maintenance_cycles mc ON t.cycle_id = mc.id
            JOIN devices d ON mc.device_id = d.id
            JOIN groups g ON d.group_id = g.id
            WHERE d.device_name IN ({placeholders})
        """
        
        params = ten_may_list
        
        if nhom and nhom != "-- Chọn nhóm --":
            query += " AND g.group_name = ?"
            params.append(nhom)
            
        print(f"[DEBUG MODEL] Câu lệnh SQL: {query}")
        print(f"[DEBUG MODEL] Tham số truyền vào: {params}")
            
        # 3. Thực thi
        try:
            with self.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute(query, params)
                
                columns = [col[0] for col in cursor.description]
                print(f"[DEBUG MODEL] Tên các cột: {columns}")
                
                results = []
                for row in cursor.fetchall():
                    # Chuyển row (tuple) thành dict
                    row_dict = dict(zip(columns, row))
                    results.append(row_dict)
                
                print(f"[DEBUG MODEL] Truy vấn thành công, trả về {len(results)} bản ghi.")
                if len(results) > 0:
                    print(f"[DEBUG MODEL] Mẫu bản ghi đầu tiên: {results[0]}")
                    print(f"[DEBUG MODEL] Kiểu dữ liệu bản ghi đầu tiên: {type(results[0])}")
                
                return results
                
        except Exception as e:
            print(f"[ERROR MODEL] Database query failed: {e}")
            import traceback
            traceback.print_exc() # In chi tiết lỗi nếu có
            return []