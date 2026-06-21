import re
from docx import Document

def parse_maintenance_word_file(file_path):
    """
    Phân tích file Word định mức bảo dưỡng.
    Tự động xử lý dấu '//', gộp phiếu (02, 03) và chuẩn hóa tên máy.
    """
    doc = Document(file_path)
    blocks = []
    
    current_device = None
    current_cycles = []
    
    # Biểu thức chính quy để bẫy Metadata
    device_pattern = re.compile(r'(?:TÊN TRANG BỊ|TÊN TRANG BỊ:)\s*(.*)', re.IGNORECASE)
    cycle_pattern = re.compile(r'PHIẾU SỐ:\s*([\d\s,]+)\s*\(([^)]+)\)', re.IGNORECASE)
    
    # Bước 1: Quét tìm các Khối Phiếu và Bảng tương ứng
    paragraphs = doc.paragraphs
    table_index = 0
    
    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        # Tìm tên máy định danh trước bảng
        dev_match = device_pattern.search(text)
        if dev_match:
            # Làm sạch tên máy (bỏ dấu ngoặc kép Nga, ký tự thừa)
            current_device = dev_match.group(1).replace('“', '').replace('”', '').strip()
            continue
            
        # Tìm số phiếu và tên chu kỳ
        cycle_match = cycle_pattern.search(text)
        if cycle_match:
            raw_nums = cycle_match.group(1)
            # Tách các số phiếu nếu ghi gộp kiểu '02, 03'
            current_cycles = [f"Phiếu {num.strip()}" for num in raw_nums.split(',') if num.strip()]
            
        # Nếu có bảng ngay sau tiêu đề, tiến hành cào dữ liệu
        if "TT" in text and "Nội dung công việc" in text and table_index < len(doc.tables):
            table = doc.tables[table_index]
            table_index += 1
            
            if not current_device or not current_cycles:
                continue
                
            tasks = []
            # Khởi tạo biến nhớ để xử lý dấu '//' (Sao chép dòng trên)
            prev_tool = ""
            prev_material = ""
            prev_tckt = ""
            
            # Duyệt qua các hàng của bảng (Bỏ qua 2 hàng tiêu đề đầu)
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                
                # Loại bỏ cấu trúc hàng trùng lặp do merged cells dọc
                if not cells or cells[0] == "TT" or "Nội dung công việc" in cells[1]:
                    continue
                    
                stt = cells[0]
                task_name = cells[1]
                
                # Bỏ dòng Tổng cộng
                if "Tổng" in stt or "Tổng" in task_name:
                    continue
                    
                # Nếu hàng trống hoàn toàn -> Bỏ qua
                if not stt and not task_name:
                    continue
                
                # Đọc các thông số định mức
                workers = cells[2] if len(cells) > 2 else ""
                minutes = cells[3] if len(cells) > 3 else ""
                tool = cells[4] if len(cells) > 4 else ""
                material = cells[5] if len(cells) > 5 else ""
                tckt = cells[6] if len(cells) > 6 else ""
                
                # --- THUẬT TOÁN ĐIỀN DỮ LIỆU ĐỘNG (LÀM SẠCH DẤU //) ---
                if tool == "//" or tool == "-":
                    tool = prev_tool if tool == "//" else ""
                elif tool:
                    prev_tool = tool
                    
                if material == "//" or material == "-":
                    material = prev_material if material == "//" else ""
                elif material:
                    prev_material = material
                    
                if tckt == "//" or tckt == "-":
                    tckt = prev_tckt if tckt == "//" else ""
                elif tckt:
                    prev_tckt = tckt
                
                # Ép kiểu dữ liệu an toàn
                try: norm_workers = int(workers)
                except: norm_workers = 0
                
                try: norm_minutes = int(minutes)
                except: norm_minutes = 0
                
                tasks.append({
                    "tt": stt,
                    "name": task_name,
                    "norm_workers": norm_workers,
                    "norm_minutes": norm_minutes,
                    "tool": tool,
                    "material": material,
                    "tckt": tckt
                })
            
            # Nhân bản khối dữ liệu vào các phiếu tương ứng (Xử lý phiếu gộp như 02, 03)
            for cycle in current_cycles:
                blocks.append({
                    "device_name": current_device,
                    "cycle_code": cycle,
                    "tasks": tasks
                })
                
    return blocks